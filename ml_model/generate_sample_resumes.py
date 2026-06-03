from docx import Document
from docx.shared import Pt, RGBColor, Inches
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from datetime import datetime

def create_docx_resume(filename, resume_data):
    """Create a DOCX resume"""
    doc = Document()

    
    title = doc.add_heading(resume_data['name'], level=0)
    title.alignment = 1  

    
    contact = doc.add_paragraph()
    contact.alignment = 1
    contact.add_run(f"{resume_data['email']} | {resume_data['phone']}").font.size = Pt(10)

    
    doc.add_heading('Professional Summary', level=2)
    doc.add_paragraph(resume_data['summary'])

    
    doc.add_heading('Technical Skills', level=2)
    skills_text = ", ".join(resume_data['skills'])
    doc.add_paragraph(skills_text)

    
    doc.add_heading('Work Experience', level=2)
    for exp in resume_data['experience']:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{exp['role']} at {exp['company']}").bold = True
        doc.add_paragraph(f"{exp['duration']}", style='List Bullet 2')

    
    doc.add_heading('Projects', level=2)
    for proj in resume_data['projects']:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(proj['name']).bold = True
        doc.add_paragraph(f"Technologies: {proj['tech']}", style='List Bullet 2')

    
    doc.add_heading('Education', level=2)
    doc.add_paragraph(f"{resume_data['education']['degree']}")
    doc.add_paragraph(f"{resume_data['education']['school']}, {resume_data['education']['year']}")

    doc.save(filename)
    print(f"[OK] Created: {filename}")

def create_pdf_resume(filename, resume_data):
    """Create a PDF resume"""
    doc = SimpleDocTemplate(filename, pagesize=letter)
    story = []
    styles = getSampleStyleSheet()

    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        textColor=colors.HexColor('#2c3e50'),
        spaceAfter=6,
        alignment=1
    )
    story.append(Paragraph(resume_data['name'], title_style))

    
    contact_style = ParagraphStyle(
        'Contact',
        parent=styles['Normal'],
        fontSize=9,
        alignment=1
    )
    story.append(Paragraph(f"{resume_data['email']} | {resume_data['phone']}", contact_style))
    story.append(Spacer(1, 0.2*inch))

    
    story.append(Paragraph('<b>Professional Summary</b>', styles['Heading2']))
    story.append(Paragraph(resume_data['summary'], styles['Normal']))
    story.append(Spacer(1, 0.1*inch))

    
    story.append(Paragraph('<b>Technical Skills</b>', styles['Heading2']))
    story.append(Paragraph(", ".join(resume_data['skills']), styles['Normal']))
    story.append(Spacer(1, 0.1*inch))

    
    story.append(Paragraph('<b>Work Experience</b>', styles['Heading2']))
    for exp in resume_data['experience']:
        story.append(Paragraph(f"<b>{exp['role']}</b> at {exp['company']}", styles['Normal']))
        story.append(Paragraph(exp['duration'], styles['Normal']))
    story.append(Spacer(1, 0.1*inch))

    
    story.append(Paragraph('<b>Projects</b>', styles['Heading2']))
    for proj in resume_data['projects']:
        story.append(Paragraph(f"<b>{proj['name']}</b>", styles['Normal']))
        story.append(Paragraph(f"Tech: {proj['tech']}", styles['Normal']))
    story.append(Spacer(1, 0.1*inch))

    
    story.append(Paragraph('<b>Education</b>', styles['Heading2']))
    story.append(Paragraph(resume_data['education']['degree'], styles['Normal']))
    story.append(Paragraph(f"{resume_data['education']['school']}, {resume_data['education']['year']}", styles['Normal']))

    doc.build(story)
    print(f"[OK] Created: {filename}")


resumes = [
    {
        'name': 'Alice Johnson',
        'email': 'alice.johnson@email.com',
        'phone': '(555) 123-4567',
        'summary': 'Passionate Full Stack Developer with 3 years of experience in building scalable web applications. Proficient in React, Node.js, and cloud technologies. Strong problem-solving skills and team player.',
        'skills': ['React', 'JavaScript', 'Node.js', 'Express', 'MongoDB', 'AWS', 'Docker', 'REST API', 'HTML', 'CSS', 'Git', 'Agile'],
        'experience': [
            {
                'role': 'Senior Frontend Developer',
                'company': 'TechCorp Solutions',
                'duration': '2022 - Present'
            },
            {
                'role': 'Full Stack Developer',
                'company': 'StartupXYZ',
                'duration': '2021 - 2022'
            },
            {
                'role': 'Junior Developer',
                'company': 'WebDev Inc',
                'duration': '2020 - 2021'
            }
        ],
        'projects': [
            {
                'name': 'E-Commerce Platform',
                'tech': 'React, Node.js, MongoDB, Stripe'
            },
            {
                'name': 'Real-time Chat Application',
                'tech': 'Socket.io, Express, PostgreSQL'
            },
            {
                'name': 'Task Management Dashboard',
                'tech': 'React, Firebase, Material-UI'
            }
        ],
        'education': {
            'degree': 'Bachelor of Science in Computer Science',
            'school': 'University of Technology',
            'year': 2020
        }
    },
    {
        'name': 'Bob Smith',
        'email': 'bob.smith@email.com',
        'phone': '(555) 234-5678',
        'summary': 'Data Science enthusiast with expertise in Python, Machine Learning, and Data Analysis. Experienced in building predictive models and data visualizations. Strong statistical knowledge.',
        'skills': ['Python', 'Machine Learning', 'TensorFlow', 'Pandas', 'NumPy', 'Scikit-learn', 'SQL', 'Data Visualization', 'Jupyter', 'Git', 'Statistics'],
        'experience': [
            {
                'role': 'Data Scientist',
                'company': 'AI Analytics Inc',
                'duration': '2021 - Present'
            },
            {
                'role': 'Data Analyst',
                'company': 'Data Solutions Ltd',
                'duration': '2019 - 2021'
            }
        ],
        'projects': [
            {
                'name': 'Customer Churn Prediction Model',
                'tech': 'Python, TensorFlow, Pandas'
            },
            {
                'name': 'Sales Forecasting System',
                'tech': 'ARIMA, Scikit-learn, SQL'
            },
            {
                'name': 'Data Dashboard',
                'tech': 'Python, Plotly, Flask'
            }
        ],
        'education': {
            'degree': 'Master of Science in Data Science',
            'school': 'State University',
            'year': 2019
        }
    },
    {
        'name': 'Carol Davis',
        'email': 'carol.davis@email.com',
        'phone': '(555) 345-6789',
        'summary': 'DevOps Engineer with 5 years of experience in cloud infrastructure, containerization, and CI/CD pipelines. Expertise in AWS, Kubernetes, and automation. Dedicated to improving deployment efficiency.',
        'skills': ['Docker', 'Kubernetes', 'AWS', 'CI/CD', 'Jenkins', 'Terraform', 'Python', 'Linux', 'Git', 'Prometheus', 'ELK Stack'],
        'experience': [
            {
                'role': 'Senior DevOps Engineer',
                'company': 'CloudTech Systems',
                'duration': '2020 - Present'
            },
            {
                'role': 'DevOps Engineer',
                'company': 'Enterprise Solutions',
                'duration': '2018 - 2020'
            },
            {
                'role': 'Systems Administrator',
                'company': 'InfoTech Services',
                'duration': '2016 - 2018'
            }
        ],
        'projects': [
            {
                'name': 'Kubernetes Migration',
                'tech': 'Docker, Kubernetes, AWS ECS'
            },
            {
                'name': 'CI/CD Pipeline Setup',
                'tech': 'Jenkins, GitLab CI, Terraform'
            },
            {
                'name': 'Infrastructure as Code',
                'tech': 'Terraform, AWS CloudFormation'
            }
        ],
        'education': {
            'degree': 'Bachelor of Science in Information Technology',
            'school': 'Technical University',
            'year': 2016
        }
    }
]

if __name__ == '__main__':
    print("Generating sample resumes...\n")

    for i, resume in enumerate(resumes, 1):
        
        docx_name = f"sample_resume_{i}.docx"
        create_docx_resume(docx_name, resume)

        
        pdf_name = f"sample_resume_{i}.pdf"
        create_pdf_resume(pdf_name, resume)

        print()

    print(f"\n[SUCCESS] Created {len(resumes)} sample resumes!")
    print("\nGenerated files:")
    for i in range(1, len(resumes) + 1):
        print(f"  • sample_resume_{i}.docx")
        print(f"  • sample_resume_{i}.pdf")
